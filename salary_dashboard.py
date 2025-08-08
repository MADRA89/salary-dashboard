import streamlit as st
import pandas as pd
import altair as alt
import random
import time
from io import BytesIO
from pandas import ExcelWriter
from docx import Document
from docx.shared import Inches

--- Helper Functions ---

def mock_parse_cv_and_jd():
return {
"educationScore": random.choice([10, 7, 4, 2]),
"experienceScore": random.choice([10, 7, 5, 3])
}

def mock_parse_interview_sheet():
return {
"performanceScore": random.choice([10, 7, 5, 2])
}

def get_step_interval(score):
if score >= 25:
return list(range(12, 16)), "Top Range"
elif score >= 20:
return list(range(9, 12)), "Mid-Upper Range"
elif score >= 15:
return list(range(6, 9)), "Mid Range"
elif score >= 10:
return list(range(3, 6)), "Lower-Mid Range"
else:
return list(range(1, 3)), "Bottom Range"

def convert_df_to_excel(df):
output = BytesIO()
with ExcelWriter(output, engine="openpyxl") as writer:
df.to_excel(writer, sheet_name="EquityAnalysis", index=False)
return output.getvalue()

def load_filtered_equity_data(uploaded_file, position_title_input):
df = pd.read_excel(uploaded_file)
df.columns = [col.strip().lower() for col in df.columns]
required_cols = {'id', 'position title', 'hire date', 'comp rate'}
if not required_cols.issubset(set(df.columns)):
return None, "❌ Excel must include columns: ID, Position Title, Hire Date, Comp Rate."
df_filtered = df[[c for c in df.columns if c in required_cols]]
df_filtered.columns = ['id', 'positionTitle', 'hireDate', 'compRate']
df_filtered = df_filtered[df_filtered['positionTitle'].str.lower() == position_title_input.lower()]
return df_filtered, None

def generate_word_report(name, title, grade, education_score, experience_score, performance_score,
total_score, interval_options, placement, selected_step,
recommended_salary, final_salary, budget_threshold, hr_comments, df_peers):
doc = Document()
doc.add_heading('Final Salary Evaluation Report', 0)

doc.add_heading('Candidate Details', level=1)  
doc.add_paragraph(f"Name: {name}")  
doc.add_paragraph(f"Position Title: {title}")  
doc.add_paragraph(f"Grade: {grade}")  

doc.add_heading('Scoring Breakdown', level=1)  
doc.add_paragraph(f"Education Score: {education_score}/10")  
doc.add_paragraph(f"Experience Score: {experience_score}/10")  
doc.add_paragraph(f"Performance Score: {performance_score}/10")  
doc.add_paragraph(f"Total Score: {total_score}/30")  
doc.add_paragraph(f"Suggested Step Interval: {interval_options} → {placement}")  
doc.add_paragraph(f"Final Selected Step: {selected_step}")  

doc.add_heading('Salary Recommendation', level=1)  
doc.add_paragraph(f"AI-Recommended Salary: AED {recommended_salary:,.0f}")  
doc.add_paragraph(f"Final Recommended Salary: AED {final_salary:,.0f}")  
doc.add_paragraph(f"Budget Threshold: AED {budget_threshold:,.0f}")  
budget_status = "Within Budget ✅" if final_salary <= budget_threshold else "Out of Budget ❌"  
doc.add_paragraph(f"Budget Status: {budget_status}")  

if df_peers is not None and not df_peers.empty:  
    doc.add_heading('Internal Equity Data', level=1)  
    table = doc.add_table(rows=1, cols=len(df_peers.columns))  
    table.style = 'Table Grid'  
    hdr_cells = table.rows[0].cells  
    for i, col_name in enumerate(df_peers.columns):  
        hdr_cells[i].text = col_name  

    for _, row in df_peers.iterrows():  
        row_cells = table.add_row().cells  
        for i, item in enumerate(row):  
            row_cells[i].text = str(item)  

doc.add_heading("HR Final Comments", level=1)  
doc.add_paragraph(hr_comments if hr_comments.strip() else "N/A")  

return doc

--- Streamlit UI ---

st.set_page_config(page_title="Salary Evaluation Dashboard", layout="wide")
st.markdown("<h1 style='color:#003366;'>📊 Salary Evaluation Dashboard</h1>", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["📘 Evaluation & Scoring Matrices", "📋 Candidate Analysis"])

with tab1:
col1, col2 = st.columns(2)
with col1:
st.subheader("🎓 Candidate Evaluation Matrix")
st.markdown("""

Criteria	Points

Master’s degree or higher	10
Bachelor’s degree	7
Diploma/Associate degree	4
High school diploma or less	2
10+ years experience	10
5-10 years experience	7
2-5 years experience	5
<2 years experience	3
High performance	10
Above average performance	7
Average performance	5
Limited performance	2


""")  

with col2:  
    st.subheader("📈 Score-to-Step Matrix")  
    st.markdown("""  
    | Score Range | Step Interval | Placement |  
    |-------------|----------------|-----------|  
    | 25–30 | Steps 12–15 | Top Range |  
    | 20–24 | Steps 9–11 | Mid-Upper Range |  
    | 15–19 | Steps 6–8 | Mid Range |  
    | 10–14 | Steps 3–5 | Lower-Mid Range |  
    | <10   | Steps 1–2 | Bottom Range |  
    """)

with tab2:
st.subheader("Step 1: Candidate & Position Details")
colA, colB = st.columns([1, 2])

with colA:  
    name = st.text_input("👤 Candidate Name")  
    title = st.text_input("🏷️ Position Title (for equity comparison)")  
    grade = st.text_input("🎖️ Position Grade")  

st.subheader("Step 2: Upload Documents")  
col1, col2, col3, col4 = st.columns(4)  
with col1:  
    uploaded_cv = st.file_uploader("📄 CV", type=["pdf", "docx"])  
with col2:  
    uploaded_jd = st.file_uploader("📝 Job Description", type=["pdf", "docx"])  
with col3:  
    uploaded_interview = st.file_uploader("🗒️ Interview Sheet", type=["pdf", "docx"])  
with col4:  
    uploaded_equity = st.file_uploader("📊 Internal Equity Excel", type=["xlsx"])  

st.subheader("Step 3: AI Evaluation + Manual Adjustment")  
if uploaded_cv and uploaded_jd:  
    with st.spinner("🔍 Evaluating CV & JD..."):  
        time.sleep(1)  
        ai_scores = mock_parse_cv_and_jd()  
else:  
    ai_scores = {"educationScore": 0, "experienceScore": 0}  

if uploaded_interview:  
    with st.spinner("🧠 Evaluating Interview..."):  
        time.sleep(1)  
        ai_scores.update(mock_parse_interview_sheet())  
else:  
    ai_scores["performanceScore"] = 0  

education_score = st.slider("🎓 Education Score (Editable)", 0, 10, ai_scores["educationScore"])  
experience_score = st.slider("💼 Experience Score (Editable)", 0, 10, ai_scores["experienceScore"])  
performance_score = st.slider("🚀 Performance Score (Editable)", 0, 10, ai_scores["performanceScore"])  

total_score = education_score + experience_score + performance_score  
interval_options, placement = get_step_interval(total_score)  

st.markdown(f"""  
### 🎯 Candidate Scoring Summary

Education: {education_score}/10 — Based on the candidate's highest qualification.

Experience: {experience_score}/10 — Reflects total relevant years of work experience.

Performance: {performance_score}/10 — Inferred from interview evaluation or performance record.

Total Score: {total_score}/30 — Combined score determining placement tier.

Suggested Step Interval: {interval_options} → {placement} — Indicates the appropriate compensation band.
""")
selected_step = st.selectbox("✅ Select Final Step", interval_options)

st.subheader("Step 4: Salary Recommendation")
budget_threshold = st.number_input("💰 Budget Threshold (AED)", step=500)
recommended_salary = st.number_input("🤖 AI-Recommended Salary (AED)", step=500)
final_salary = st.number_input("✅ Final Recommended Salary (AED)", step=500)

st.subheader("Step 5: Internal Equity Analysis")
if uploaded_equity and title and final_salary == 0:
st.info("ℹ️ Please enter the Final Recommended Salary (AED) above to perform the equity analysis.")
if uploaded_equity and title and final_salary > 0:
df_peers, error = load_filtered_equity_data(uploaded_equity, title)
if error:
st.error(error)
elif df_peers.empty:
st.warning(f"No matching peers found for position title: '{title}'")
else:
df_peers = pd.concat([
df_peers,
pd.DataFrame([{
"id": "Candidate",
"positionTitle": title,
"hireDate": "N/A",
"compRate": final_salary
}])
], ignore_index=True)

st.dataframe(df_peers.style.set_properties(**{'background-color': '#f0f0f5'}))  

      avg = df_peers["compRate"].mean()  
      min_val = df_peers["compRate"].min()  
      max_val = df_peers["compRate"].max()  

      st.markdown(f"**Equity Range for '{title}':** Min AED {min_val:,} | Avg AED {avg:,.0f} | Max AED {max_val:,}")  

      chart = alt.Chart(df_peers).mark_bar().encode(  
          x=alt.X("id:N", title="Employee ID"),  
          y=alt.Y("compRate:Q", title="Compensation (AED)"),  
          color=alt.condition(alt.datum.id == "Candidate", alt.value("#FF8C00"), alt.value("#2a5c88")),  
          tooltip=["id", "compRate"]  
      ).properties(title="💼 Compensation Comparison", width=700, height=350)  

      st.altair_chart(chart, use_container_width=True)  

      excel_data = convert_df_to_excel(df_peers)  
      st.download_button(  
          label="📥 Download Equity Data (.xlsx)",  
          data=excel_data,  
          file_name="equity_analysis_filtered.xlsx",  
          mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"  
      )

st.subheader("Step 6: Budget Status")
if final_salary and budget_threshold:
if final_salary <= budget_threshold:
st.success(f"✅ Within Budget (AED {final_salary:,.0f})")
else:
st.error(f"❌ Out of Budget (AED {final_salary:,.0f})")

st.subheader("Step 7: Final HR Comments & Export")
hr_comments = st.text_area("📝 HR Final Comments")

if st.button("📤 Generate Final Summary"):
summary = f"""
📌 Final Recommendation Summary

Candidate Name: {name}  
  Position Title: {title}  
  Grade: {grade}  

  Total Score: {total_score}/30 → Step Interval: {interval_options} → Placement: {placement}  
  Selected Step: {selected_step}  
  AI-Recommended Salary: AED {recommended_salary:,.0f}  
  Final Recommended Salary: AED {final_salary:,.0f}  
  Budget Threshold: AED {budget_threshold:,.0f}  
  Budget Status: {'Within' if final_salary <= budget_threshold else 'Out of'} Budget  

  HR Final Comments:  
  {hr_comments}  
  """  
  st.text_area("📋 Final Summary", summary, height=250)  
  st.download_button("📤 Download Final Summary (.txt)", data=summary, file_name="salary_summary.txt")  

  # Generate Word document  
  doc = generate_word_report(  
      name, title, grade,  
      education_score, experience_score, performance_score,  
      total_score, interval_options, placement, selected_step,  
      recommended_salary, final_salary, budget_threshold,  
      hr_comments, df_peers if 'df_peers' in locals() else None  
  )  

  buffer = BytesIO()  
  doc.save(buffer)  
  buffer.seek(0)  

  st.download_button(  
      label="📥 Download Full Report (.docx)",  
      data=buffer,  
      file_name=f"{name.replace(' ', '_').lower()}_evaluation_report.docx",  
      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"  
  )


